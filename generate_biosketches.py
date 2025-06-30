#!/usr/bin/env python3
"""
Generate biosketches and core usage analysis for CDTR members.
This script uses data from UAB Scholars API to create comprehensive biosketches
and analyze core usage for each faculty member.
"""

import csv
import json
import re
from datetime import datetime
from typing import Dict, List, Any, Optional
import pandas as pd
from cdtr_collaboration_pull import (
    fetch_user_js, fetch_all_pages, clean_text, find_user_id,
    API_BASE, USERS_API_BASE, USERS_API_SEARCH,
    PUBS_API_URL, GRANTS_API_URL, TEACH_API_URL,
    HEADERS, PER_PAGE_PUBS, PER_PAGE_GRANTS, PER_PAGE_TEACHING,
    MANUAL_DISCOVERY_IDS
)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import scholars_api_shim  # noqa: F401

# Constants
CDTR_MEMBERS_CSV = "cdtr members.csv"
CDTR_COLLABORATIONS_CSV = "cdtr_collaborations.csv"
OUTPUT_DIR = "biosketches"

def extract_credentials(positions: List[Dict]) -> str:
    """Extract credentials from positions."""
    credentials = set()
    for pos in positions:
        if "credentials" in pos:
            creds = pos["credentials"].split(",")
            credentials.update(c.strip() for c in creds)
    return ", ".join(sorted(credentials))

def extract_roles(positions: List[Dict]) -> Dict[str, Any]:
    """Extract primary and secondary roles from positions."""
    roles = []
    for pos in positions:
        role = {
            "role": pos.get("position", ""),
            "division": pos.get("division", ""),
            "department": pos.get("department", ""),
            "school": pos.get("school", ""),
            "university": pos.get("university", ""),
            "city": pos.get("city", ""),
            "state": pos.get("state", "")
        }
        roles.append(role)
    
    return {
        "primary": roles[0] if roles else {},
        "secondary": roles[1:] if len(roles) > 1 else []
    }

def is_diabetes_related(grant: Dict[str, Any]) -> bool:
    """Check if a grant is diabetes-related."""
    diabetes_keywords = [
        "diabetes", "diabetic", "glucose", "insulin", "metabolic",
        "obesity", "weight", "glycemic", "HbA1c", "A1C"
    ]
    
    title = grant.get("title", "").lower()
    description = grant.get("description", "").lower()
    labels = "; ".join(l.get("value", "") for l in grant.get("labels", [])).lower()
    
    # Debugging print statements
    print(f"Grant Title: {title}")
    print(f"Grant Description: {description}")
    print(f"Grant Labels: {labels}")
    
    text = f"{title} {description} {labels}"
    return any(keyword in text for keyword in diabetes_keywords)

def get_recent_diabetes_grants(grants: List[Dict[str, Any]], limit: int = 10) -> List[Dict[str, Any]]:
    """Get the most recent diabetes-related grants."""
    diabetes_grants = [g for g in grants if is_diabetes_related(g)]
    return sorted(
        diabetes_grants,
        key=lambda x: x.get("date1", {}).get("dateTime", ""),
        reverse=True
    )[:limit]

def generate_research_overview(profile: Dict[str, Any], publications: List[Dict[str, Any]]) -> str:
    """Generate research overview narrative."""
    bio = profile.get("bio", "")
    research_interests = profile.get("researchInterests", "")
    
    # Extract key points from bio and research interests
    points = []
    if bio:
        points.append(bio)
    if research_interests:
        points.append(f"Research interests include {research_interests}")
    
    # Add publication count
    pub_count = len(publications)
    points.append(f"Has published {pub_count} scholarly works")
    
    return " ".join(points)

def parse_core_descriptions() -> Dict[str, Any]:
    """Parse core descriptions from the provided documents."""
    # This is a placeholder - we'll need to implement proper document parsing
    # For now, returning hardcoded core descriptions
    return {
        "translational_design": {
            "name": "Translational Design & Intervention Core",
            "description": """
            The Translational Design & Intervention Core provides expertise in:
            - Community-engaged research methods
            - Intervention development and testing
            - Implementation science
            - Health disparities research
            - Patient-reported outcomes
            - Practice-based research
            """,
            "keywords": [
                "community", "intervention", "implementation", "disparities",
                "patient", "practice", "clinical", "trial", "randomized",
                "qualitative", "quantitative", "mixed methods"
            ]
        },
        "data_science": {
            "name": "Data Science & Analytics Core",
            "description": """
            The Data Science & Analytics Core provides expertise in:
            - Big data analytics
            - Machine learning
            - Statistical analysis
            - Data visualization
            - Electronic health records
            - Real-world data
            """,
            "keywords": [
                "data", "analytics", "statistics", "machine learning",
                "visualization", "EHR", "electronic health records",
                "big data", "predictive", "modeling"
            ]
        }
    }

def analyze_core_usage(profile: Dict[str, Any], grants: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Analyze how the faculty member will use CDTR cores."""
    core_descriptions = parse_core_descriptions()
    
    # Combine all text for analysis
    text = " ".join([
        profile.get("bio", ""),
        profile.get("researchInterests", ""),
        profile.get("teachingSummary", ""),
        " ".join(g.get("title", "") for g in grants),
        " ".join(g.get("description", "") for g in grants)
    ]).lower()
    
    # Analyze core usage
    core_usage = {}
    for core_id, core in core_descriptions.items():
        # Check if research aligns with core keywords
        keyword_matches = sum(1 for k in core["keywords"] if k in text)
        will_use = keyword_matches >= 2  # At least 2 keyword matches
        
        # Generate usage description
        if will_use:
            matched_keywords = [k for k in core["keywords"] if k in text]
            usage = f"Will use the {core['name']} for {', '.join(matched_keywords)}-related research and analysis."
        else:
            usage = f"May use the {core['name']} for specific project needs."
        
        core_usage[core_id] = {
            "will_use": will_use,
            "usage": usage,
            "matched_keywords": matched_keywords if will_use else []
        }
    
    return core_usage

def process_faculty_member(name: str) -> Dict[str, Any]:
    """Process a single faculty member's data."""
    # Try manual override first
    disc_id = MANUAL_DISCOVERY_IDS.get(name)
    if not disc_id:
        # Convert "Last, First Middle" to "First Middle Last"
        if ',' in name:
            last, first = name.split(',', 1)
            search_name = f"{first.strip()} {last.strip()}"
        else:
            search_name = name
        disc_id = find_user_id(search_name)
    if not disc_id:
        print(f"Could not find user ID for {name}")
        return None
    js = fetch_user_js(disc_id)
    if not js:
        print(f"Error: Could not fetch user {disc_id} for {name}")
        return None
    
    # Enhanced profile information
    profile = {
        "objectId": js.get("objectId", ""),
        "discoveryUrlId": js.get("discoveryUrlId", ""),
        "firstName": js.get("firstName", ""),
        "lastName": js.get("lastName", ""),
        "email": js.get("email", ""),
        "orcid": js.get("orcid", ""),
        "bio": clean_text(js.get("bio", "")),
        "researchInterests": clean_text(js.get("researchInterests", "")),
        "teachingSummary": clean_text(js.get("teachingSummary", "")),
        "positions": js.get("positions", []),
        "credentials": extract_credentials(js.get("positions", [])),
        "roles": extract_roles(js.get("positions", [])),
        "education": js.get("education", []),
        "awards": js.get("awards", []),
        "professionalActivities": js.get("professionalActivities", [])
    }
    
    # Fetch publications with enhanced details
    publications = []
    for page in fetch_all_pages(
        PUBS_API_URL,
        lambda s: {
            "objectId": disc_id,
            "objectType": "user",
            "pagination": {"perPage": PER_PAGE_PUBS, "startFrom": s},
            "favouritesFirst": True,
            "sort": "dateDesc"
        },
        PER_PAGE_PUBS
    ):
        for pub in page:
            enhanced_pub = {
                **pub,
                "citation": pub.get("citation", ""),
                "journal": pub.get("journal", ""),
                "publicationDate": pub.get("publicationDate", ""),
                "doi": pub.get("doi", ""),
                "abstract": pub.get("abstract", "")
            }
            publications.append(enhanced_pub)
    
    # Fetch grants with enhanced details
    grants = []
    for page in fetch_all_pages(
        GRANTS_API_URL,
        lambda s: {
            "objectId": disc_id,
            "objectType": "user",
            "pagination": {"perPage": PER_PAGE_GRANTS, "startFrom": s},
            "favouritesFirst": True,
            "sort": "dateDesc"
        },
        PER_PAGE_GRANTS
    ):
        for grant in page:
            enhanced_grant = {
                **grant,
                "description": grant.get("description", ""),
                "amount": grant.get("amount", ""),
                "startDate": grant.get("date1", {}).get("dateTime", ""),
                "endDate": grant.get("date2", {}).get("dateTime", ""),
                "status": grant.get("status", ""),
                "role": grant.get("role", ""),
                "labels": grant.get("labels", [])
            }
            grants.append(enhanced_grant)
    
    # Generate outputs
    return {
        "profile": profile,
        "research_overview": generate_research_overview(profile, publications),
        "core_usage": analyze_core_usage(profile, grants),
        "recent_diabetes_grants": get_recent_diabetes_grants(grants),
        "publications": publications,
        "grants": grants,
        "education": profile.get("education", []),
        "awards": profile.get("awards", []),
        "professional_activities": profile.get("professionalActivities", [])
    }

def load_json_data(faculty_name):
    """Load all JSON data for a faculty member."""
    data = {}
    json_files = {
        'profile': f'faculty_profiles/{faculty_name}.json',
        'publications': f'faculty_publications/{faculty_name}.json',
        'grants': f'faculty_grants/{faculty_name}.json',
        'teaching': f'faculty_teaching/{faculty_name}.json'
    }
    
    for key, file_path in json_files.items():
        try:
            with open(file_path, 'r') as f:
                data[key] = json.load(f)
        except FileNotFoundError:
            data[key] = None
    
    return data

def generate_personal_statement(profile_data):
    """Generate a personal statement based on profile data."""
    if not profile_data:
        return "No profile data available."
    
    statement = []
    
    # Add research focus
    if 'research_focus' in profile_data:
        statement.append(f"My research focuses on {profile_data['research_focus']}.")
    
    # Add key achievements
    if 'key_achievements' in profile_data:
        statement.append(f"Key achievements include {profile_data['key_achievements']}.")
    
    # Add current work
    if 'current_work' in profile_data:
        statement.append(f"Currently, I am working on {profile_data['current_work']}.")
    
    return " ".join(statement)

def format_positions_and_honors(profile_data):
    """Format positions and honors chronologically."""
    if not profile_data or 'positions' not in profile_data:
        return "No position data available."
    
    positions = []
    for pos in profile_data['positions']:
        date = pos.get('date', '')
        title = pos.get('title', '')
        institution = pos.get('institution', '')
        positions.append(f"{date}: {title}, {institution}")
    
    return "\n".join(positions)

def format_publications(pub_data):
    """Format recent publications."""
    if not pub_data or 'publications' not in pub_data:
        return "No publication data available."
    
    # Sort by date and get most recent
    pubs = sorted(pub_data['publications'], 
                 key=lambda x: x.get('date', ''), 
                 reverse=True)[:5]
    
    formatted_pubs = []
    for pub in pubs:
        authors = pub.get('authors', '')
        title = pub.get('title', '')
        journal = pub.get('journal', '')
        date = pub.get('date', '')
        formatted_pubs.append(f"{authors}. {title}. {journal}. {date}.")
    
    return "\n".join(formatted_pubs)

def format_research_support(grant_data):
    """Format active and pending grants."""
    if not grant_data or 'grants' not in grant_data:
        return "No grant data available."
    
    # Filter for active and pending grants
    active_grants = [g for g in grant_data['grants'] 
                    if g.get('status', '').lower() in ['active', 'pending']]
    
    formatted_grants = []
    for grant in active_grants:
        title = grant.get('title', '')
        agency = grant.get('agency', '')
        amount = grant.get('amount', '')
        period = grant.get('period', '')
        role = grant.get('role', '')
        formatted_grants.append(
            f"{title}\n{agency}\nAmount: {amount}\nPeriod: {period}\nRole: {role}\n"
        )
    
    return "\n".join(formatted_grants)

def create_biosketch(faculty_name):
    """Create a biosketch document for a faculty member."""
    # Load data
    data = load_json_data(faculty_name)
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Biosketch: {faculty_name}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add sections
    doc.add_heading('Personal Statement', level=2)
    doc.add_paragraph(generate_personal_statement(data['profile']))
    
    doc.add_heading('Positions and Honors', level=2)
    doc.add_paragraph(format_positions_and_honors(data['profile']))
    
    doc.add_heading('Selected Peer-reviewed Publications', level=2)
    doc.add_paragraph(format_publications(data['publications']))
    
    doc.add_heading('Research Support', level=2)
    doc.add_paragraph(format_research_support(data['grants']))
    
    # Save document
    output_dir = 'biosketches'
    os.makedirs(output_dir, exist_ok=True)
    doc.save(f'{output_dir}/{faculty_name} biosketch.docx')

def main():
    """Generate biosketch for Gareth Dutton only."""
    faculty_name = "Gareth Dutton"
    print(f"Generating biosketch for {faculty_name}...")
    
    # Process the faculty member
    data = process_faculty_member(faculty_name)
    if data:
        # Create output directory if it doesn't exist
        os.makedirs("biosketches", exist_ok=True)
        
        # Generate timestamp for filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = f"biosketches/biosketch_{faculty_name.replace(' ', '_')}_{timestamp}.json"
        
        # Save as JSON
        with open(output_file, 'w') as f:
            json.dump(data, f, indent=2)
        print(f"Completed biosketch for {faculty_name}")
        print(f"Output saved to: {output_file}")
    else:
        print(f"Failed to generate biosketch for {faculty_name}")

if __name__ == "__main__":
    main() 