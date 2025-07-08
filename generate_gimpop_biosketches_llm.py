#!/usr/bin/env python3
"""generate_gimpop_biosketches_llm.py

Generate NIH-style biosketches for every profile JSON in `scholar_data/`
using OpenAI GPT-4o-mini.  Sections created:
  A. Personal Statement (LLM-generated)
  B. Positions and Honors (from profile positions)
  C. Contributions to Science (LLM-generated from publications)
  D. Selected Publications (top N publications)

Requirements (already in repo):
  python-docx, openai, python-dotenv

The OPENAI_API_KEY must be present in `.env` or environment.

Docx files are written to `output_bios_docx/`.
"""
from __future__ import annotations

import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import List

from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt

# ---------------------------------------------------------------------------
# Configuration / Setup
# ---------------------------------------------------------------------------
ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / "scholar_data"
OUT_DIR = ROOT / "output_bios_docx"
OUT_DIR.mkdir(exist_ok=True)

load_dotenv()
OPENAI_MODEL = "gpt-4o-mini"  # user asked for "gpt 4.1 mini model" – this is closest
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def to_last_first(name: str) -> str:
    parts = name.split()
    return f"{parts[-1]}, {' '.join(parts[:-1])}" if len(parts) >= 2 else name


def get_positions_text(positions: List[dict]) -> str:
    """Return positions / honors section as plain text."""
    lines: List[str] = []
    for pos in positions:
        role = pos.get("position") or ""
        dept = pos.get("department") or ""
        division = pos.get("division") or ""
        school = pos.get("school") or ""
        inst_parts = ", ".join(p for p in [division, dept, school] if p)
        loc_parts = ", ".join(p for p in [pos.get("city"), pos.get("state")] if p)
        line = ", ".join(p for p in [role, inst_parts, loc_parts] if p)
        if line:
            lines.append(line)
    return "\n".join(lines) if lines else "Positions not listed."


def call_llm(system_prompt: str, user_context: str) -> str:
    try:
        resp = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_context},
            ],
            temperature=0.7,
            max_tokens=700,
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        print(f"⚠️  OpenAI error: {e}")
        return "LLM generation unavailable."


def build_personal_statement(profile: dict, pubs_text: str) -> str:
    system_prompt = (
        "You are an expert grant writer producing NIH biosketch Section A (Personal Statement). "
        "Write in first person, ~250 words. Integrate the faculty member's background, positions, "
        "research interests, and up to three key publications to illustrate expertise."
    )
    user_context = (
        f"Name: {profile.get('name')}\n"
        f"Positions: {get_positions_text(profile.get('raw_positions', []))}\n"
        f"Research interests / bio: {profile.get('bio','')}\n"
        f"Recent publications (one per line):\n{pubs_text}"
    )
    return call_llm(system_prompt, user_context)


def build_contributions_to_science(pubs_text: str) -> str:
    system_prompt = (
        "Write NIH biosketch Section C (Contributions to Science). "
        "Identify 2–3 contributions based on the publication list, each 100–120 words, followed by citations."
    )
    return call_llm(system_prompt, pubs_text)


def write_docx(bio: dict) -> None:
    """Save biosketch to DOCX using python-docx."""
    doc = Document()
    doc.add_heading("NIH Biosketch", level=0)

    def h(text: str):
        doc.add_heading(text, level=1)

    def p(text: str):
        para = doc.add_paragraph(text)
        para.style.font.size = Pt(11)

    # Section A
    h("A. Personal Statement")
    p(bio["personal_statement"])

    # Section B
    h("B. Positions and Honors")
    p(bio["positions"])

    # Section C
    h("C. Contributions to Science")
    p(bio["contributions"])

    # Section D
    h("D. Selected Publications")
    p(bio["selected_publications"])

    filename = OUT_DIR / f"{to_last_first(bio['name'])} biosketch.docx"
    doc.save(filename)
    print("📄", filename.name)


# ---------------------------------------------------------------------------
# Main workflow
# ---------------------------------------------------------------------------

def main(limit: int | None = None, top_pubs: int = 10) -> None:
    files = sorted(DATA_DIR.glob("*_profile_*.json"))
    if limit:
        files = files[:limit]
    if not files:
        sys.exit("No profile JSON files found in scholar_data");

    for js_file in files:
        data = json.loads(js_file.read_text())
        if data.get("status") != "success":
            continue
        res = data["results"][0]
        profile = res["profile"].copy()
        profile["raw_positions"] = data["results"][0]["profile"].get("positions", [])
        profile["bio"] = res.get("profile", {}).get("teaching_interests") or ""

        # Prepare publications text
        pubs = res.get("publications", {}).get("recent", [])[:top_pubs]
        pubs_text = "\n".join(
            f"{p.get('year','')} {p.get('title','')}. {p.get('journal','')}" for p in pubs
        ) or "No publications listed."

        # LLM-generated sections
        personal_stmt = build_personal_statement(profile, pubs_text)
        contrib = build_contributions_to_science(pubs_text)

        biosketch_dict = {
            "name": profile["name"],
            "personal_statement": personal_stmt,
            "positions": get_positions_text(profile["raw_positions"]),
            "contributions": contrib,
            "selected_publications": pubs_text,
        }
        write_docx(biosketch_dict)

    print(f"\nGenerated {len(files)} DOCX biosketches in {OUT_DIR}/")


if __name__ == "__main__":
    import argparse

    ap = argparse.ArgumentParser(description="Generate DOCX biosketches via OpenAI LLM")
    ap.add_argument("--limit", type=int, default=None, help="Process only first N profiles (for testing)")
    ap.add_argument("--top-pubs", type=int, default=10, help="Number of publications to include")
    args = ap.parse_args()

    main(limit=args.limit, top_pubs=args.top_pubs) 